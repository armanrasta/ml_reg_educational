import pandas as pd
import numpy as np
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error
from datetime import datetime
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
import os

class ModelReportGenerator:
    def __init__(self, output_dir='reports'):
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
    def generate_report(self, model, X_train, X_test, y_train, y_test, y_pred, feature_importance):
        doc = Document()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc.add_heading('Car Price Prediction Model Report', 0)
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Model Information
        doc.add_heading('Model Information', level=1)
        doc.add_paragraph(f'Model Type: {type(model).__name__}')
        doc.add_paragraph(f'Number of features: {X_train.shape[1]}')
        doc.add_paragraph(f'Training samples: {X_train.shape[0]}')
        doc.add_paragraph(f'Test samples: {X_test.shape[0]}')
        
        # Model Performance
        doc.add_heading('Model Performance Metrics', level=1)
        r2 = r2_score(y_test, y_pred)
        mse = mean_squared_error(y_test, y_pred)
        rmse = np.sqrt(mse)
        mae = mean_absolute_error(y_test, y_pred)
        
        metrics = [
            ('R-squared Score', r2),
            ('Mean Squared Error', mse),
            ('Root Mean Squared Error', rmse),
            ('Mean Absolute Error', mae)
        ]
        
        for metric_name, value in metrics:
            doc.add_paragraph(f'{metric_name}: {value:.4f}')
        
        # Feature Importance
        if feature_importance is not None:
            self._add_feature_importance(doc, feature_importance)
        
        # Model Visualizations
        self._add_visualizations(doc, y_test, y_pred, X_train)
        
        # Save report
        report_path = os.path.join(self.output_dir, f'model_report_{timestamp}.docx')
        doc.save(report_path)
        return report_path

    def _add_feature_importance(self, doc, feature_importance):
        doc.add_heading('Feature Importance Analysis', level=1)
        plt.figure(figsize=(10, 6))
        plt.barh(feature_importance['Feature'], feature_importance['Importance'])
        plt.title('Feature Importance')
        plt.tight_layout()
        img_path = os.path.join(self.output_dir, 'feature_importance.png')
        plt.savefig(img_path)
        doc.add_picture(img_path, width=Inches(6))
        plt.close()

    def _add_visualizations(self, doc, y_test, y_pred, X_train):
        doc.add_heading('Model Visualizations', level=1)
        
        # Actual vs Predicted
        plt.figure(figsize=(10, 6))
        plt.scatter(y_test, y_pred, alpha=0.5)
        plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], 'r--')
        plt.xlabel('Actual Price')
        plt.ylabel('Predicted Price')
        plt.title('Actual vs Predicted Prices')
        plt.tight_layout()
        img_path = os.path.join(self.output_dir, 'actual_vs_predicted.png')
        plt.savefig(img_path)
        doc.add_picture(img_path, width=Inches(6))
        plt.close()
        
        # Correlation Matrix
        try:
            X_train_df = pd.DataFrame(X_train)
            plt.figure(figsize=(12, 8))
            sns.heatmap(X_train_df.corr(), annot=True, cmap='coolwarm')
            plt.title('Feature Correlation Matrix')
            plt.tight_layout()
            img_path = os.path.join(self.output_dir, 'correlation_matrix.png')
            plt.savefig(img_path)
            doc.add_picture(img_path, width=Inches(6))
            plt.close()
        except Exception as e:
            doc.add_paragraph(f"Could not generate correlation matrix: {str(e)}")
